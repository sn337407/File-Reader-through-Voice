"""Any extension File Reader through voice"""

import pyttsx #module to convert text into speech
import docx #module to read the document
import winsound #module to play audio
from pyPdf import PdfFileReader #module to convert a pdf into text

""" Text to speech"""

engine=pyttsx.init() 
engine.setProperty('rate', 180) #setting up voice rate
x = raw_input("Enter the File") # I am using a file named sample.txt as input from the python folder
input_file= x + ".txt"
file= open(input_file,"rt") #open a file "sample.txt"
test_in_file=file.read()
engine.say(test_in_file) #Reading a file through voice
engine.runAndWait()

winsound.PlaySound('sound.wav',winsound.SND_FILENAME) #palying an audio indicates an end of file
print ("TTS Success")

"""PDF to Speech"""

y = raw_input("Enter the pdf") # I am using a file name sample.pdf as input from the pyhton folder
input_file = y + ".pdf"
pdf=open(input_file,'rb') #Reading a pdf
reader = PdfFileReader(pdf) #converting a pdf to text
content = reader.getPage(0).extractText().split('\n') #extracting the texts from pdf
engine.say(content)
engine.runAndWait()

winsound.PlaySound('sound.wav',winsound.SND_FILENAME)
print('PTS Success')

"""Document to Speech"""

engine.setProperty('rate', 100)
engine.setProperty('volume', 0.5)
voices = engine.getProperty('voices') 
z = raw_input("Enter the Document") #I am using a file named sample.docx as input  from python folder
input_file= z + ".docx"
for voice in voices: #to change the voice of speech
    engine.setProperty('voice', voice.id)
doc = docx.Document(input_file) #opening a document
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
read=getText(input_file) #Reading a document
engine.say(read)
engine.runAndWait()

winsound.PlaySound('sound.wav',winsound.SND_FILENAME)
print('DTS Success')

